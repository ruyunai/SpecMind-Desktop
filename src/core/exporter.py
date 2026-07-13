"""PRD 结果导出模块。

支持格式：
- Markdown：适合版本管理、人工审阅
- Word (.docx)：适合交付客户/内部评审
- JSON：完整 State，适合审计与后续系统对接
"""
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def _now_str() -> str:
    """获取当前时间字符串。"""
    return datetime.now().strftime("%Y-%m-%d %H:%M")


def _extract_weeks(phase: Dict[str, Any]) -> int:
    """从交付阶段提取周数，兼容 duration/weeks key 与 int/str 类型。

    Args:
        phase: 交付阶段字典，可能含 duration 或 weeks 字段

    Returns:
        周数（int），解析失败返回 0
    """
    for key in ("duration", "weeks"):
        val = phase.get(key)
        if val is None:
            continue
        if isinstance(val, (int, float)):
            return int(val)
        if isinstance(val, str):
            # 提取首个数字（兼容 "3周"/"约4周"/"1.5周"）
            import re
            match = re.search(r"[\d.]+", val)
            if match:
                try:
                    return int(float(match.group()))
                except ValueError:
                    return 0
            return 0
    return 0


def _format_duration(phase: Dict[str, Any]) -> str:
    """格式化阶段工期为展示字符串，兼容 duration/weeks key 与 int/str 类型。

    Args:
        phase: 交付阶段字典

    Returns:
        工期展示字符串（如 "3周" / "0周"）
    """
    for key in ("duration", "weeks"):
        val = phase.get(key)
        if val is None:
            continue
        if isinstance(val, (int, float)):
            return f"{int(val)}周"
        if isinstance(val, str) and val.strip():
            return val
    return "0周"


def export_prd_as_markdown(state: Dict[str, Any], client_name: str = "") -> str:
    """将 PRD 及配套结果导出为 Markdown 字符串。

    Args:
        state: 工作流最终 State
        client_name: 客户名

    Returns:
        Markdown 格式文本
    """
    lines = []
    title_client = f" - {client_name}" if client_name else ""
    lines.append(f"# SpecMind PRD 输出{title_client}")
    lines.append(f"\n生成时间：{_now_str()}\n")

    # 1. SAR 清洗报告
    cleaned = state.get("cleaned_requirements", "")
    risks = state.get("overcommit_risks", [])
    if cleaned or risks:
        lines.append("## 一、SAR 需求清洗报告")
        lines.append(f"\n### 清洗后需求\n\n{cleaned}\n")
        if risks:
            lines.append("### 过度承诺风险\n")
            for i, risk in enumerate(risks, 1):
                lines.append(f"{i}. {risk}")
            lines.append("")

    # 2. PRD 8 模块
    prd = state.get("prd", {})
    if prd:
        lines.append("## 二、PRD 文档（8 模块）\n")
        for module, content in prd.items():
            lines.append(f"### {module}")
            lines.append(f"\n{content}\n")

    # 3. 功能点标注
    features = state.get("prd_features", [])
    if features:
        lines.append("## 三、功能点标注\n")
        lines.append("| 功能 | 标注 | 说明 |")
        lines.append("|------|------|------|")
        for feat in features:
            name = feat.get("name", "").replace("|", "\\|")
            tag = feat.get("tag", "").replace("|", "\\|")
            desc = feat.get("desc", "").replace("|", "\\|")
            lines.append(f"| {name} | {tag} | {desc} |")
        lines.append("")

    # 4. 报价
    quotes = state.get("quotes", {})
    if quotes:
        lines.append("## 四、Commercial 双报价\n")
        for version, quote in quotes.items():
            lines.append(f"### {version}")
            for k, v in quote.items():
                if isinstance(v, float) and v < 1:
                    lines.append(f"- {k}: {v:.0%}")
                else:
                    lines.append(f"- {k}: {v}")
            lines.append("")

    # 5. 合同冲突
    conflicts = state.get("contract_conflicts", [])
    if conflicts:
        lines.append("## 五、Contract 合同冲突\n")
        for i, c in enumerate(conflicts, 1):
            lines.append(f"### 冲突 {i} [{c.get('risk', '').upper()}]")
            lines.append(f"- PRD 条款：{c.get('prd_clause', '')}")
            lines.append(f"- 合同条款：{c.get('contract_clause', '')}")
            lines.append(f"- 冲突说明：{c.get('conflict', '')}")
            lines.append("")

    # 6. 评审意见
    reviews = state.get("review_comments", {})
    if reviews:
        lines.append("## 六、Review 多维评审\n")
        for dimension, comments in reviews.items():
            lines.append(f"### {dimension.upper()}")
            for c in comments:
                lines.append(f"- {c}")
            lines.append("")
        lines.append(f"**评审结论：** {'通过' if state.get('review_pass') else '不通过'}\n")

    # 7. 交付计划
    plan = state.get("delivery_plan", [])
    if plan:
        lines.append("## 七、Planner 交付计划\n")
        total_weeks = 0
        for phase in plan:
            total_weeks += _extract_weeks(phase)
            duration_str = _format_duration(phase)
            lines.append(f"### {phase.get('phase', '')}（{duration_str}）")
            lines.append(f"- 交付物：{phase.get('deliverable', phase.get('deliverables', ''))}\n")
        lines.append(f"**总工期：{total_weeks} 周**\n")

    # 8. Legal 预检声明
    lines.append("---\n")
    lines.append(
        "> ⚠ **Legal Agent 为辅助预检工具，输出不构成正式法律意见，"
        "需由正式法务人员复核后方可作为决策依据。**\n"
    )

    return "\n".join(lines)


def export_prd_as_word(state: Dict[str, Any], output_path: Path, client_name: str = "") -> None:
    """将 PRD 及配套结果导出为 Word 文档。

    Args:
        state: 工作流最终 State
        output_path: 输出文件路径
        client_name: 客户名
    """
    doc = Document()

    # 标题
    title = doc.add_heading("SpecMind PRD 输出", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if client_name:
        doc.add_paragraph(f"客户：{client_name}", style="Subtitle")
    doc.add_paragraph(f"生成时间：{_now_str()}", style="Subtitle")

    # SAR 清洗报告
    cleaned = state.get("cleaned_requirements", "")
    risks = state.get("overcommit_risks", [])
    if cleaned or risks:
        doc.add_heading("一、SAR 需求清洗报告", level=1)
        if cleaned:
            doc.add_heading("清洗后需求", level=2)
            doc.add_paragraph(cleaned)
        if risks:
            doc.add_heading("过度承诺风险", level=2)
            for risk in risks:
                doc.add_paragraph(risk, style="List Bullet")

    # PRD 8 模块
    prd = state.get("prd", {})
    if prd:
        doc.add_heading("二、PRD 文档（8 模块）", level=1)
        for module, content in prd.items():
            doc.add_heading(module, level=2)
            doc.add_paragraph(content)

    # 功能点标注
    features = state.get("prd_features", [])
    if features:
        doc.add_heading("三、功能点标注", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "功能"
        hdr_cells[1].text = "标注"
        hdr_cells[2].text = "说明"
        for feat in features:
            row_cells = table.add_row().cells
            row_cells[0].text = str(feat.get("name", ""))
            row_cells[1].text = str(feat.get("tag", ""))
            row_cells[2].text = str(feat.get("desc", ""))

    # 报价
    quotes = state.get("quotes", {})
    if quotes:
        doc.add_heading("四、Commercial 双报价", level=1)
        for version, quote in quotes.items():
            doc.add_heading(version, level=2)
            for k, v in quote.items():
                if isinstance(v, float) and v < 1:
                    doc.add_paragraph(f"{k}: {v:.0%}", style="List Bullet")
                else:
                    doc.add_paragraph(f"{k}: {v}", style="List Bullet")

    # 合同冲突
    conflicts = state.get("contract_conflicts", [])
    if conflicts:
        doc.add_heading("五、Contract 合同冲突", level=1)
        for i, c in enumerate(conflicts, 1):
            doc.add_heading(f"冲突 {i} [{c.get('risk', '').upper()}]", level=2)
            doc.add_paragraph(f"PRD 条款：{c.get('prd_clause', '')}")
            doc.add_paragraph(f"合同条款：{c.get('contract_clause', '')}")
            doc.add_paragraph(f"冲突说明：{c.get('conflict', '')}")

    # 评审意见
    reviews = state.get("review_comments", {})
    if reviews:
        doc.add_heading("六、Review 多维评审", level=1)
        for dimension, comments in reviews.items():
            doc.add_heading(dimension.upper(), level=2)
            for c in comments:
                doc.add_paragraph(c, style="List Bullet")
        doc.add_paragraph(f"评审结论：{'通过' if state.get('review_pass') else '不通过'}")

    # 交付计划
    plan = state.get("delivery_plan", [])
    if plan:
        doc.add_heading("七、Planner 交付计划", level=1)
        total_weeks = 0
        for phase in plan:
            total_weeks += _extract_weeks(phase)
            duration_str = _format_duration(phase)
            doc.add_heading(f"{phase.get('phase', '')}（{duration_str}）", level=2)
            doc.add_paragraph(f"交付物：{phase.get('deliverable', phase.get('deliverables', ''))}")
        doc.add_paragraph(f"总工期：{total_weeks} 周")

    # Legal 声明
    doc.add_paragraph()
    legal_para = doc.add_paragraph(
        "Legal Agent 为辅助预检工具，输出不构成正式法律意见，"
        "需由正式法务人员复核后方可作为决策依据。"
    )
    legal_para.runs[0].font.italic = True
    legal_para.runs[0].font.size = Pt(9)

    doc.save(str(output_path))


def export_full_report_as_json(state: Dict[str, Any], output_path: Path) -> None:
    """将完整 State 导出为 JSON 文件（审计/对接用途）。

    Args:
        state: 工作流最终 State
        output_path: 输出文件路径
    """
    report = {
        "export_time": _now_str(),
        "state": state,
    }
    output_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
